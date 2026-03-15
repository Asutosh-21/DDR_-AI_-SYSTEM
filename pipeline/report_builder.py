import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from config import OUTPUT_DOCX, OUTPUT_DIR


def set_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0] if heading.runs else heading.add_run(text)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return heading


def add_section_label(doc, label, value):
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(11)
    p.add_run(str(value)).font.size = Pt(11)


def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_images(doc, image_paths, label=""):
    valid = [p for p in image_paths if p and os.path.exists(p)]
    if not valid:
        doc.add_paragraph(f"{label}Image Not Available").italic = True
        return
    if label:
        doc.add_paragraph(label).bold = True
    for path in valid:
        try:
            doc.add_picture(path, width=Inches(2.8))
        except Exception:
            doc.add_paragraph(f"[Could not load image: {os.path.basename(path)}]")
    doc.add_paragraph()


def build_report(ddr):
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    title = doc.add_heading("Detailed Diagnostic Report (DDR)", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    doc.add_paragraph()

    # 1. Property Issue Summary
    set_heading(doc, "1. Property Issue Summary")
    doc.add_paragraph(ddr.get("property_issue_summary", "Not Available"))
    add_divider(doc)

    # 2. Area-wise Observations
    set_heading(doc, "2. Area-wise Observations")
    for i, obs in enumerate(ddr.get("area_wise_observations", []), 1):
        set_heading(doc, f"  Area {i}: {obs.get('area', 'Unknown')}", level=2)
        add_section_label(doc, "Issue Observed", obs.get("negative_side", "Not Available"))
        add_section_label(doc, "Source / Cause Area", obs.get("positive_side", "Not Available"))
        add_section_label(doc, "Thermal Finding", obs.get("thermal_finding", "Not Available"))

        # Inspection images
        insp_imgs = obs.get("inspection_images", [])
        add_images(doc, insp_imgs, "Inspection Photos: ")

        # Thermal images
        thermal_imgs = obs.get("thermal_images", [])
        add_images(doc, thermal_imgs, "Thermal Images: ")

        doc.add_paragraph()

    add_divider(doc)

    # 3. Probable Root Cause
    set_heading(doc, "3. Probable Root Cause")
    doc.add_paragraph(ddr.get("probable_root_cause", "Not Available"))
    add_divider(doc)

    # 4. Severity Assessment
    set_heading(doc, "4. Severity Assessment")
    severity = ddr.get("severity_assessment", {})
    add_section_label(doc, "Severity Level", severity.get("level", "Not Available"))
    add_section_label(doc, "Reasoning", severity.get("reasoning", "Not Available"))
    add_divider(doc)

    # 5. Recommended Actions
    set_heading(doc, "5. Recommended Actions")
    for action in ddr.get("recommended_actions", ["Not Available"]):
        doc.add_paragraph(action, style="List Bullet")
    add_divider(doc)

    # 6. Additional Notes
    set_heading(doc, "6. Additional Notes")
    doc.add_paragraph(ddr.get("additional_notes", "Not Available"))
    add_divider(doc)

    # 7. Missing or Unclear Information
    set_heading(doc, "7. Missing or Unclear Information")
    missing = ddr.get("missing_or_unclear_info", [])
    if missing:
        for item in missing:
            doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph("Not Available")

    doc.save(OUTPUT_DOCX)
    print(f"DDR Report saved to: {OUTPUT_DOCX}")


if __name__ == "__main__":
    import json
    from config import EXTRACTED_DIR
    with open(f"{EXTRACTED_DIR}/ddr_with_images.json", "r", encoding="utf-8") as f:
        ddr = json.load(f)
    build_report(ddr)
